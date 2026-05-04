import os
from docx.shared import Inches
from converter.strategies.base_strategy import BaseStrategy

class HybridStrategy(BaseStrategy):

    def process(self, page, doc, index):
        if self.is_complex(page, index):
            self._add_as_image(page, doc, index)
        else:
            self._add_as_text(page, doc, index)

    def is_complex(self, page, index=None):
        text = page.get_text("text").strip()
        images = page.get_images()
        drawings = page.get_drawings()

        text_len = len(text)
        images_len = len(images)
        drawings_len = len(drawings)

        print("\n---------------------------")
        print(f"Página {index}")
        print(f"Tamanho do texto: {text_len}")
        print(f"Imagens: {images_len}")
        print(f"Drawings: {drawings_len}")

        # 🔥 NOVA LÓGICA

        # Pouquíssimo texto + imagem → provavelmente scan/layout
        if text_len < 30 and images_len > 0:
            print("➡️ COMPLEXO (pouco texto + imagem)")
            return True

        # Muitas imagens → complexo
        if images_len > 3:
            print("➡️ COMPLEXO (muitas imagens)")
            return True

        # Layout muito pesado
        if drawings_len > 50:
            print("➡️ COMPLEXO (muitos drawings)")
            return True

        print("➡️ SIMPLES (texto)")
        return False

    def _add_as_text(self, page, doc, index):
        data = page.get_text("dict")

        elements = []

        for block in data["blocks"]:
            # TEXTO
            if block["type"] == 0:
                text = ""
                for line in block["lines"]:
                    for span in line["spans"]:
                        text += span["text"]

                if text.strip():
                    elements.append(("text", block["bbox"][1], text))

            # IMAGEM
            elif block["type"] == 1:
                bbox = block["bbox"]
                y = bbox[1]

                image_bytes = block["image"]
                path = f"temp_img_{index}_{y}.png"

                with open(path, "wb") as f:
                    f.write(image_bytes)

                elements.append(("image", y, path))

        # 🔥 Ordena por posição vertical
        elements.sort(key=lambda x: x[1])

        for element in elements:
            if element[0] == "text":
                doc.add_paragraph(element[2])
            else:
                doc.add_picture(element[2], width=Inches(4))
                os.remove(element[2])

    def _add_as_image(self, page, doc, index):
        pix = page.get_pixmap()
        path = f"temp_page_{index}.png"
        pix.save(path)

        doc.add_picture(path, width=Inches(6))

        os.remove(path)